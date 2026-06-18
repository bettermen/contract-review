#!/usr/bin/env python3
"""
合同文本解析工具 — 从 docx/pdf/txt 文件中提取文本内容。
用法: python parse_contract.py <file_path> [--type auto|docx|pdf|txt]
"""

import sys
import argparse
from pathlib import Path


def extract_txt(filepath: str) -> str:
    """Extract text from .txt file."""
    with open(filepath, 'r', encoding='utf-8') as f:
        return f.read()


def extract_docx(filepath: str) -> str:
    """Extract text from .docx file using python-docx."""
    try:
        from docx import Document
    except ImportError:
        sys.exit("ERROR: python-docx not installed. Run: pip install python-docx")

    doc = Document(filepath)
    paragraphs = []
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text.strip())

    # Also extract tables
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(' | '.join(cells))

    return '\n\n'.join(paragraphs)


def extract_pdf(filepath: str) -> str:
    """Extract text from .pdf file."""
    try:
        import pdfplumber
    except ImportError:
        try:
            from PyPDF2 import PdfReader
        except ImportError:
            sys.exit("ERROR: Need pdfplumber or PyPDF2. Run: pip install pdfplumber")

    # Try pdfplumber first (better text extraction)
    try:
        import pdfplumber
        with pdfplumber.open(filepath) as pdf:
            pages = []
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    pages.append(text.strip())
            return '\n\n'.join(pages)
    except ImportError:
        pass

    # Fallback to PyPDF2
    from PyPDF2 import PdfReader
    reader = PdfReader(filepath)
    pages = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            pages.append(text.strip())
    return '\n\n'.join(pages)


def identify_contract_type(text: str) -> str:
    """Identify contract type from keywords."""
    text_lower = text.lower()

    keywords = {
        '劳动用工合同': ['劳动合同', '劳动用工', '聘用', '雇佣', '试用期', '工资', '社保', '甲方（用人单位'],
        '买卖合同': ['买卖合同', '购销', '采购', '供货', '货物', '交货', '验收', '标的物'],
        '租赁合同': ['租赁', '出租', '承租', '租金', '押金', '租赁物', '转租'],
        '服务合同': ['服务合同', '技术服务', '咨询服务', '委托服务', '服务内容', '服务费'],
        '保密协议': ['保密协议', '保密', '秘密', '竞业', 'NDA', '商业秘密', '保密信息'],
        '投资协议': ['投资', '股权', '增资', '估值', '赎回', '优先清算', '反稀释', '领售'],
        '建设工程合同': ['建设工程', '施工', '承包', '竣工', '工程质量', '工期', '农民工'],
        '技术开发合同': ['技术开发', '研发', '开发', '源代码', '软件', '知识产权归属', '里程碑'],
    }

    scores = {}
    for ctype, kws in keywords.items():
        scores[ctype] = sum(1 for kw in kws if kw in text_lower)

    if not scores or max(scores.values()) == 0:
        return '通用合同'

    return max(scores, key=scores.get)


def main():
    parser = argparse.ArgumentParser(description='合同文本解析工具')
    parser.add_argument('filepath', help='合同文件路径 (.txt/.docx/.pdf)')
    parser.add_argument('--type', choices=['auto', 'docx', 'pdf', 'txt'], default='auto',
                        help='文件类型（默认自动检测）')
    parser.add_argument('--identify', action='store_true', help='尝试识别合同类型')

    args = parser.parse_args()
    filepath = Path(args.filepath)

    if not filepath.exists():
        sys.exit(f"ERROR: File not found: {args.filepath}")

    # Determine file type
    if args.type == 'auto':
        suffix = filepath.suffix.lower()
        if suffix == '.docx':
            file_type = 'docx'
        elif suffix == '.pdf':
            file_type = 'pdf'
        else:
            file_type = 'txt'
    else:
        file_type = args.type

    # Extract text
    extractors = {
        'txt': extract_txt,
        'docx': extract_docx,
        'pdf': extract_pdf,
    }

    text = extractors[file_type](str(filepath))

    print(f"=== 合同文本提取成功 ({file_type.upper()}) ===")
    print(f"=== 文件: {filepath.name} ===")
    print(f"=== 字符数: {len(text)} ===")

    if args.identify:
        ctype = identify_contract_type(text)
        print(f"=== 识别合同类型: {ctype} ===")

    print("\n--- 合同内容 ---\n")
    print(text)


if __name__ == '__main__':
    main()
